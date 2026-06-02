"""
INGESTION SERVICE
Handles: File reading → Text extraction → Chunking into smaller pieces

Why chunking?
   An AI model has a limited context window. You can't send an entire
   200-page PDF at once. Instead, we split the text into small chunks
   (e.g., 500 characters each), embed each chunk, and later retrieve
   only the RELEVANT chunks for a given question.
"""

import os
import re
from typing import List
from PyPDF2 import PdfReader
from docx import Document
from app.core.config import settings


def ensure_upload_dir():
    """Create the uploads directory if it doesn't exist."""
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file."""
    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file."""
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def extract_text_from_txt(file_path: str) -> str:
    """Read a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(file_path: str, filename: str) -> str:
    """Extract text from a file based on its extension."""
    extension = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if extension == "pdf":
        return extract_text_from_pdf(file_path)
    elif extension == "docx":
        return extract_text_from_docx(file_path)
    elif extension in ("txt", "md", "csv"):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: .{extension}")


def clean_text(text: str) -> str:
    """Clean extracted text — remove excessive whitespace, normalize."""
    # Replace multiple newlines with double newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Replace multiple spaces with single space
    text = re.sub(r" {2,}", " ", text)
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    return text.strip()


def chunk_text(text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[str]:
    """
    Split text into overlapping chunks at sentence boundaries.

    Why sentence boundaries?
       If we cut mid-sentence, the chunk loses meaning.
       "Our pizza is made with fresh" ← useless
       "Our pizza is made with fresh mozzarella and basil." ← useful

    Why overlap?
       If a relevant answer spans two chunks, the overlap ensures
       context isn't lost at the boundary.
    """
    if chunk_size is None:
        chunk_size = settings.CHUNK_SIZE
    if chunk_overlap is None:
        chunk_overlap = settings.CHUNK_OVERLAP

    if not text or len(text.strip()) == 0:
        return []

    text = clean_text(text)

    # If the entire text fits in one chunk, return as-is
    if len(text) <= chunk_size:
        return [text]

    # Split into sentences using punctuation boundaries
    sentences = re.split(r"(?<=[.!?])\s+", text)

    chunks = []
    current_chunk = ""

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        # If adding this sentence exceeds the chunk size
        if len(current_chunk) + len(sentence) + 1 > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())

            # If a single sentence is longer than chunk_size, force-split it
            if len(sentence) > chunk_size:
                words = sentence.split()
                current_chunk = ""
                for word in words:
                    if len(current_chunk) + len(word) + 1 > chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = word
                    else:
                        current_chunk += " " + word if current_chunk else word
            else:
                current_chunk = sentence
        else:
            current_chunk += " " + sentence if current_chunk else sentence

    # Don't forget the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    # Apply overlap — prepend the tail of the previous chunk to the next
    if chunk_overlap > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev = chunks[i - 1]
            overlap_text = prev[-chunk_overlap:] if len(prev) > chunk_overlap else prev
            # Find word boundary in overlap to avoid cutting words
            space_idx = overlap_text.find(" ")
            if space_idx != -1:
                overlap_text = overlap_text[space_idx + 1:]
            overlapped.append(overlap_text + " " + chunks[i])
        chunks = overlapped

    # Filter out chunks that are too short to be useful
    chunks = [c for c in chunks if len(c) > 20]

    return chunks